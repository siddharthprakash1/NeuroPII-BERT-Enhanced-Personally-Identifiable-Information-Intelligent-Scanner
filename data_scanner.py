'''
import re
import pandas as pd
import torch
from transformers import AutoTokenizer, AutoModelForTokenClassification
from docx import Document
from flask import Flask, request, render_template, redirect, url_for
import os

class SensitiveDataScanner:
    def __init__(self):
        # Load BERT model and tokenizer
        self.tokenizer = AutoTokenizer.from_pretrained('dslim/bert-base-NER')
        self.model = AutoModelForTokenClassification.from_pretrained('dslim/bert-base-NER')
        
        # Enhanced regex patterns with additional medical record formats
        self.patterns = {
            'pan_numbers': r'[A-Z]{5}[0-9]{4}[A-Z]',  # ABCDE1234F format
            'ssn': r'\d{3}-\d{2}-\d{4}',  # 123-45-6789 format
            'medical_record': r'(?:MR\d{6}|\d{3}-\d{2}-\d{4}(?!\s*(?:SSN|Social Security))|MRN\d+)',  # Multiple medical record formats
            'credit_card': r'(?:\d{4}[-\s]?){4}'  # Various credit card formats
        }

    def scan_text(self, text):
        results = {
            'pan_numbers': set(),
            'ssn': set(),
            'medical_record': set(),
            'credit_card': set()
        }
        
        # First find SSNs that are explicitly labeled as SSN
        ssn_with_label = re.finditer(r'(?i)(?:SSN|Social Security).*?(\d{3}-\d{2}-\d{4})', str(text))
        explicit_ssns = {match.group(1) for match in ssn_with_label}
        
        # Search for patterns in the text
        for pattern_type, pattern in self.patterns.items():
            matches = re.findall(pattern, str(text))
            # For medical records, exclude numbers that were identified as explicit SSNs
            if pattern_type == 'medical_record':
                matches = [m for m in matches if m not in explicit_ssns]
            results[pattern_type].update(matches)
            
        # Add explicit SSNs to SSN results
        results['ssn'].update(explicit_ssns)
            
        return results

def extract_text_from_docx(file):
    """Extract text content from a Word document."""
    doc = Document(file)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def scan_file(file):
    scanner = SensitiveDataScanner()
    
    try:
        # Check file extension
        if file.filename.endswith('.csv'):
            # Read CSV file
            df = pd.read_csv(file)
            text_content = df.to_string()
        elif file.filename.endswith(('.xls', '.xlsx')):
            # Read Excel file
            df = pd.read_excel(file)
            text_content = df.to_string()
        elif file.filename.endswith(('.doc', '.docx')):
            # Read Word document
            text_content = extract_text_from_docx(file)
        else:
            raise ValueError("Unsupported file format. Please use CSV, Excel, or Word files.")
        
        # Scan for sensitive data
        results = scanner.scan_text(text_content)
        
        # Format results
        scan_result = {
            'filename': file.filename,
            'pan_numbers': ', '.join(results['pan_numbers']) if results['pan_numbers'] else 'No PAN card numbers found',
            'ssn_numbers': ', '.join(results['ssn']) if results['ssn'] else 'No SSN numbers found',
            'medical_record_numbers': ', '.join(results['medical_record']) if results['medical_record'] else 'No medical record numbers found',
            'credit_card_numbers': ', '.join(results['credit_card']) if results['credit_card'] else 'No credit card numbers found'
        }
        
        # Print detailed results for debugging
        print("\nScanning file:", file.filename)
        print("\nDetected Sensitive Data:")
        for category, items in results.items():
            print(f"\n{category.replace('_', ' ').title()}:")
            if items:
                for item in items:
                    print(f"- {item}")
            else:
                print("None found")
        
        return scan_result
        
    except Exception as e:
        print(f"Error processing file: {str(e)}")
        return {
            'filename': file.filename,
            'pan_numbers': 'Error processing file',
            'ssn_numbers': 'Error processing file',
            'medical_record_numbers': 'Error processing file',
            'credit_card_numbers': 'Error processing file'
        }

# Flask application setup
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

scans = []

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return redirect(request.url)
    
    file = request.files['file']
    if file.filename == '':
        return redirect(request.url)
    
    if file and (file.filename.endswith('.csv') or 
                 file.filename.endswith('.xlsx') or 
                 file.filename.endswith('.xls') or
                 file.filename.endswith('.doc') or
                 file.filename.endswith('.docx')):
        # Process the file
        result = scan_file(file)
        scans.append(result)
        return redirect(url_for('scan_results'))
    
    return 'Invalid file type'

@app.route('/scan_results')
def scan_results():
    return render_template('scan_results.html', scans=scans)

@app.route('/delete_scan/<int:scan_id>')
def delete_scan_route(scan_id):
    if 0 <= scan_id < len(scans):
        scans.pop(scan_id)
    return redirect(url_for('scan_results'))

if __name__ == '__main__':
    app.run(debug=True)
'''
import re
import pandas as pd
import torch
from transformers import AutoTokenizer, AutoModelForTokenClassification
from docx import Document

class SensitiveDataScanner:
    def __init__(self):
        # Load BERT model and tokenizer
        self.tokenizer = AutoTokenizer.from_pretrained('dslim/bert-base-NER')
        self.model = AutoModelForTokenClassification.from_pretrained('dslim/bert-base-NER')
        
        # Enhanced regex patterns
        self.patterns = {
            'pan_numbers': r'[A-Z]{5}[0-9]{4}[A-Z]',
            'ssn': r'\d{3}-\d{2}-\d{4}',
            'medical_record': r'(?:MR\d{6}|\d{3}-\d{2}-\d{4}(?!\s*(?:SSN|Social Security))|MRN\d+)',
            'credit_card': r'(?:\d{4}[-\s]?){4}'
        }

    def scan_text(self, text):
        results = {
            'pan_numbers': set(),
            'ssn': set(),
            'medical_record': set(),
            'credit_card': set()
        }
        
        # First find SSNs that are explicitly labeled as SSN
        ssn_with_label = re.finditer(r'(?i)(?:SSN|Social Security).*?(\d{3}-\d{2}-\d{4})', str(text))
        explicit_ssns = {match.group(1) for match in ssn_with_label}
        
        # Search for patterns in the text
        for pattern_type, pattern in self.patterns.items():
            matches = re.findall(pattern, str(text))
            if pattern_type == 'medical_record':
                matches = [m for m in matches if m not in explicit_ssns]
            results[pattern_type].update(matches)
            
        results['ssn'].update(explicit_ssns)
        return results

def extract_text_from_docx(file):
    """Extract text content from a Word document."""
    try:
        doc = Document(file)
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text from docx: {str(e)}")
        return ""

def scan_file(file):
    scanner = SensitiveDataScanner()
    
    try:
        # Check file extension
        if not hasattr(file, 'filename'):
            raise ValueError("Invalid file object")
            
        if file.filename.endswith('.csv'):
            df = pd.read_csv(file)
            text_content = df.to_string()
        elif file.filename.endswith(('.xls', '.xlsx')):
            df = pd.read_excel(file)
            text_content = df.to_string()
        elif file.filename.endswith(('.doc', '.docx')):
            # Reset file pointer to beginning
            file.seek(0)
            text_content = extract_text_from_docx(file)
        else:
            result = {
                'filename': getattr(file, 'filename', 'unknown'),
                'pan_numbers': 'Error processing file',
                'ssn_numbers': 'Error processing file',
                'medical_record_numbers': 'Error processing file',
                'credit_card_numbers': 'Error processing file'
            }
            print(f"Error processing file: Unsupported file format. Please use CSV, Excel, or Word files.")
            return result
        
        # If we got here, we have text content to scan
        if not text_content:
            results = {
                'pan_numbers': set(),
                'ssn': set(),
                'medical_record': set(),
                'credit_card': set()
            }
        else:
            results = scanner.scan_text(text_content)
        
        scan_result = {
            'filename': file.filename,
            'pan_numbers': ', '.join(results['pan_numbers']) if results['pan_numbers'] else 'No PAN card numbers found',
            'ssn_numbers': ', '.join(results['ssn']) if results['ssn'] else 'No SSN numbers found',
            'medical_record_numbers': ', '.join(results['medical_record']) if results['medical_record'] else 'No medical record numbers found',
            'credit_card_numbers': ', '.join(results['credit_card']) if results['credit_card'] else 'No credit card numbers found'
        }
        
        return scan_result
        
    except Exception as e:
        print(f"Error processing file: {str(e)}")
        return {
            'filename': getattr(file, 'filename', 'unknown'),
            'pan_numbers': 'Error processing file',
            'ssn_numbers': 'Error processing file',
            'medical_record_numbers': 'Error processing file',
            'credit_card_numbers': 'Error processing file'
        }