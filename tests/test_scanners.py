import pytest
import os
import sys
from io import BytesIO, SEEK_SET, SEEK_CUR, SEEK_END

# Add the parent directory to Python path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from data_scanner import SensitiveDataScanner, scan_file, extract_text_from_docx
from database import store_scan_result, get_all_scans, delete_scan
from docx import Document

class MockFile:
    def __init__(self, filename, content):
        self.filename = filename
        self._content = BytesIO(content)
        
    def read(self, size=-1):
        """Read at most size bytes from the file."""
        return self._content.read(size)
        
    def seek(self, offset, whence=SEEK_SET):
        """Seek to specified position."""
        return self._content.seek(offset, whence)
        
    def tell(self):
        """Return current position."""
        return self._content.tell()
    
    def close(self):
        """Close the file."""
        self._content.close()

@pytest.fixture
def sample_docx():
    # Create a proper Word document
    doc = Document()
    doc.add_paragraph('SSN: 123-45-6789')
    doc.add_paragraph('PAN: ABCDE1234F')
    doc.add_paragraph('Credit Card: 1234-5678-9012-3456')
    doc.add_paragraph('Medical Record: MR123456')
    
    # Save to BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    content = file_stream.getvalue()
    file_stream.close()
    
    # Return mock file with the content
    return MockFile('test.docx', content)

@pytest.fixture
def scanner():
    return SensitiveDataScanner()

def test_scanner_initialization(scanner):
    assert scanner.tokenizer is not None
    assert scanner.model is not None
    assert len(scanner.patterns) == 4

def test_pan_number_detection(scanner):
    text = "PAN number is ABCDE1234F"
    results = scanner.scan_text(text)
    assert 'ABCDE1234F' in results['pan_numbers']

def test_ssn_detection(scanner):
    text = "SSN is 123-45-6789"
    results = scanner.scan_text(text)
    assert '123-45-6789' in results['ssn']

def test_credit_card_detection(scanner):
    text = "CC: 1234-5678-9012-3456"
    results = scanner.scan_text(text)
    assert '1234-5678-9012-3456' in results['credit_card']

def test_medical_record_detection(scanner):
    text = "MR: MR123456"
    results = scanner.scan_text(text)
    assert 'MR123456' in results['medical_record']

def test_docx_extraction(sample_docx):
    text = extract_text_from_docx(sample_docx)
    assert '123-45-6789' in text
    assert 'ABCDE1234F' in text

def test_scan_file_docx(sample_docx):
    result = scan_file(sample_docx)
    assert 'ABCDE1234F' in result['pan_numbers']
    assert '123-45-6789' in result['ssn_numbers']

def test_database_operations():
    # Test storing a scan
    result = {'filename': 'test.txt', 'pan_numbers': 'ABCDE1234F'}
    scan_id = store_scan_result(result)
    
    # Test retrieving scans
    scans = get_all_scans()
    assert len(scans) > 0
    assert scans[-1]['filename'] == 'test.txt'
    
    # Test deleting a scan
    delete_scan(scan_id)
    scans = get_all_scans()
    assert not any(scan.get('id') == scan_id for scan in scans)

def test_invalid_file():
    invalid_file = MockFile('test.txt', b'invalid content')
    result = scan_file(invalid_file)
    assert 'Error processing file' in result['pan_numbers']

def test_empty_file():
    # Create an empty docx file with minimal valid structure
    doc = Document()
    file_stream = BytesIO()
    doc.save(file_stream)
    content = file_stream.getvalue()
    file_stream.close()
    
    empty_file = MockFile('test.docx', content)
    result = scan_file(empty_file)
    assert 'No PAN card numbers found' in result['pan_numbers']